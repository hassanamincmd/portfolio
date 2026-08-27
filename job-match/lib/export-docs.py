"""Write resume/cover letter to DOCX and PDF from JSON on stdin."""

import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.colors import HexColor
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

INK = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x5A, 0x5A, 0x5A)


def set_run(run, size=11, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color or INK


def write_docx(text: str, path: Path, title: str):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    first = True
    for raw in text.splitlines():
        line = raw.rstrip()
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.15
        if not line:
            p.paragraph_format.space_after = Pt(6)
            continue
        upper = line.isupper() and len(line) > 3
        heading = line in {
            "PROFESSIONAL SUMMARY",
            "SKILLS",
            "PROFESSIONAL EXPERIENCE",
            "ADDITIONAL EXPERIENCE",
            "EDUCATION",
            "CERTIFICATIONS",
            "LANGUAGES",
            "PORTFOLIO",
        }
        if first:
            set_run(p.add_run(line), size=16, bold=True)
            first = False
        elif heading or upper:
            set_run(p.add_run(line), size=11, bold=True)
        elif line.startswith("- "):
            p.style = "List Bullet"
            set_run(p.add_run(line[2:]), size=10.5)
        else:
            set_run(p.add_run(line), size=10.5)
    doc.save(path)


def write_pdf(text: str, path: Path):
    doc = SimpleDocTemplate(
        str(path),
        pagesize=letter,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
        topMargin=0.65 * inch,
        bottomMargin=0.65 * inch,
    )
    base = getSampleStyleSheet()["Normal"]
    body = ParagraphStyle(
        "Body",
        parent=base,
        fontName="Helvetica",
        fontSize=10,
        leading=13,
        textColor=HexColor("#1a1a1a"),
        spaceAfter=3,
        alignment=TA_LEFT,
    )
    head = ParagraphStyle(
        "Head",
        parent=body,
        fontName="Helvetica-Bold",
        fontSize=11,
        spaceBefore=8,
        spaceAfter=4,
    )
    name = ParagraphStyle(
        "Name",
        parent=body,
        fontName="Helvetica-Bold",
        fontSize=15,
        spaceAfter=6,
    )
    story = []
    first = True
    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            story.append(Spacer(1, 6))
            continue
        safe = (
            line.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
        )
        heading = line in {
            "PROFESSIONAL SUMMARY",
            "SKILLS",
            "PROFESSIONAL EXPERIENCE",
            "ADDITIONAL EXPERIENCE",
            "EDUCATION",
            "CERTIFICATIONS",
            "LANGUAGES",
            "PORTFOLIO",
        }
        if first:
            story.append(Paragraph(safe, name))
            first = False
        elif heading or (line.isupper() and len(line) > 3):
            story.append(Paragraph(safe, head))
        else:
            story.append(Paragraph(safe, body))
    doc.build(story)


def main():
    payload = json.load(sys.stdin)
    path = Path(payload["path"])
    path.parent.mkdir(parents=True, exist_ok=True)
    kind = payload["format"]
    text = payload["text"]
    title = payload.get("title", "Document")
    if kind == "docx":
        write_docx(text, path, title)
    elif kind == "pdf":
        write_pdf(text, path)
    else:
        raise SystemExit(f"unknown format {kind}")
    print(path)


if __name__ == "__main__":
    main()
