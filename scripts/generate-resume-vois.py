"""ATS resume tailored to VOIS Junior UI/UX Designer posting."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.colors import HexColor
from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import HRFlowable, Paragraph, SimpleDocTemplate, Spacer

OUT_DIR = Path(r"C:\Users\Ryzen Store\Desktop\CV Resume")
BASE = "Hassan-Amin-CV-Resume-VOIS-UIUX"

INK = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x5A, 0x5A, 0x5A)
INK_HEX = "#1A1A1A"
MUTED_HEX = "#5A5A5A"
RULE_HEX = "#C8C8C8"

TARGET_TITLE = "UI/UX Designer"

CONTACT_LINES = [
    "Cairo, Egypt | +20 10 668 74 777 | contact.hassan.amin@gmail.com",
    "LinkedIn: linkedin.com/in/hassan-mo-amin | Portfolio: https://hassanamin.net",
]

SUMMARY = (
    "UI/UX Designer with 5 years delivering screen-level web and mobile interfaces, "
    "from wireframes and prototypes to developer-ready UI. Strong in Figma component "
    "libraries, design systems, responsive layout, accessibility (WCAG), and "
    "design-to-development handoff. Experience supporting usability research and "
    "first-pass synthesis, and using AI tools (Figma, Cursor, Claude) to speed "
    "prototyping and production alignment. Former VOIS UK Account Advisor / SME "
    "(2018-2021). Open to UI/UX design delivery in Egypt."
)

SKILLS = [
    (
        "UI / UX Delivery",
        "Wireframes, User Flows, Prototypes, High-Fidelity UI, Interaction States, "
        "Visual Hierarchy, Layout and Spacing, Responsive Design, Web Platforms, "
        "On-Brand UI, Developer-Ready Handoff, Annotations, Specifications, Asset Exports",
    ),
    (
        "Design Systems and Tools",
        "Figma, FigJam, Adobe XD, Photoshop, Illustrator, Framer, Zeplin; "
        "Component Libraries, Auto-Layout, Variants, Variables, Design Tokens, "
        "Instance Management, Production Alignment",
    ),
    (
        "Research and Methods",
        "User Research, Usability Testing, UserTesting, Note-Taking, First-Pass Synthesis, "
        "UX Principles, Accessibility (WCAG), Agile, Design Thinking",
    ),
    (
        "Collaboration and AI",
        "Jira, Notion, Trello, Stakeholder Communication, Engineering Handoff; "
        "Cursor, Claude, ChatGPT, AI-Assisted Prototyping",
    ),
]

EXPERIENCE = [
    {
        "title": "Product Designer 2",
        "company": "Procore Technologies",
        "dates": "May 2025 - July 2026",
        "meta": "Full-time | Quality and Safety | Web and mobile",
        "bullets": [
            "Designed screen-level user interfaces for 3 cross-platform Quality and Safety products, covering layout, hierarchy, interaction states, and responsive behaviour.",
            "Prepared developer-ready files aligned to the shared design system; contributed reusable components and flagged design vs production gaps.",
            "Supported usability research and first-pass synthesis, including a UserTesting study that informed redesign of 4 of 5 dashboard cards before handoff.",
            "Used Figma plus Cursor and Claude to accelerate prototyping and design-to-development handoff while keeping WCAG and brand standards.",
        ],
    },
    {
        "title": "Senior UI/UX Designer",
        "company": "Caspian Digital Solutions",
        "dates": "November 2022 - February 2025",
        "meta": "Full-time | Agile product delivery",
        "bullets": [
            "Delivered wireframes, prototypes, and high-fidelity UI across 20+ digital products per year on agreed timelines.",
            "Maintained design guidelines and component consistency; improved handoff quality with engineering and stakeholders.",
            "Presented and explained design decisions with evidence; 95% stakeholder approval and 100% on-time delivery.",
        ],
    },
    {
        "title": "UI/UX Designer",
        "company": "Caspian Digital Solutions",
        "dates": "May 2022 - October 2022",
        "meta": "Full-time",
        "bullets": [
            "Created wireframes and user flows that improved cross-team clarity during agile delivery.",
            "Established design guidelines that reduced iteration time by 25% and strengthened brand consistency.",
            "Prepared designs for review and handoff; 90% first-round approval.",
        ],
    },
    {
        "title": "UI/UX Designer / Product Designer",
        "company": "Spiritude",
        "dates": "December 2021 - May 2022",
        "meta": "Full-time",
        "bullets": [
            "Translated requirements into simple web and product interfaces using wireframes, personas, and usage scenarios.",
            "Supported research through competitor analysis and user research; delivered design-system foundations that reduced UI inconsistency.",
        ],
    },
    {
        "title": "UK Account Advisor / Subject Matter Expert",
        "company": "VOIS (Vodafone Intelligent Solutions)",
        "dates": "March 2018 - March 2021",
        "meta": "Full-time | Customer experience operations",
        "bullets": [
            "Handled UK customer escalations and documented workflows across the Customer Experience Cycle.",
            "Mentored newcomers and acted as a knowledge source for frontline advisors on policies and difficult cases.",
            "Helped improve CX processes and service journeys, which later informed how I design usable product flows.",
        ],
    },
]

FREELANCE = [
    {
        "title": "Senior Product Designer",
        "company": "Deloitte",
        "dates": "2025",
        "meta": "Freelance",
        "bullets": [
            "Extended the KSA Capital Market Authority design system with scalable UI components, variants, and library constraints.",
        ],
    },
    {
        "title": "Senior Product Designer / Design Lead",
        "company": "Path2Live",
        "dates": "March 2024 - December 2024",
        "meta": "Consultant | Concurrent with full-time role",
        "bullets": [
            "Delivered responsive web, tablet, and mobile UI from a token-based design system (100+ components) with clear engineering handoff.",
            "Documented edge cases and specifications that reduced design-engineering back-and-forth.",
        ],
    },
]

EDUCATION = [
    (
        "Intensive Training Program - UI/UX Design",
        "Information Technology Institute (ITI), Ministry of Communications and Information Technology | April 2021 - July 2021",
    ),
    (
        "Bachelor of Arts - Geographic Information Systems",
        "Faculty of Arts, King Marriott Academy | October 2013 - August 2018",
    ),
]

CERTIFICATIONS = [
    "UX Design Professional - Google",
    "Digital Product Management - Institute of Management, Technology and Finance",
    "Graphic Design Practice License - Ministry of Manpower, Egypt",
    "PMP Training Course - Arab Academy for Science, Technology and Maritime Transport",
]


def set_run_font(run, size=10.5, bold=False, color=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color if color is not None else INK


def set_paragraph_spacing(p, before=0, after=0, line=1.15):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line


def add_bottom_border(paragraph, color="C8C8C8", size="12"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def section_heading(doc, text):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=11, after=5, line=1.0)
    run = p.add_run(text.upper())
    set_run_font(run, size=10, bold=True, color=INK)
    add_bottom_border(p)
    return p


def body_para(doc, text, size=10, after=6, before=0, justify=False, color=None):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=before, after=after, line=1.15)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_run_font(run, size=size, color=color)
    return p


def bullet_para(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_paragraph_spacing(p, before=0, after=2, line=1.15)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.first_line_indent = Inches(-0.15)
    if p.runs:
        p.runs[0].text = text
        set_run_font(p.runs[0], size=10)
    else:
        run = p.add_run(text)
        set_run_font(run, size=10)
    return p


def add_job_block(doc, title, company, dates, meta=None, bullets=None):
    title_p = doc.add_paragraph()
    set_paragraph_spacing(title_p, before=8, after=0, line=1.1)
    run = title_p.add_run(title)
    set_run_font(run, size=10.5, bold=True)

    company_p = doc.add_paragraph()
    set_paragraph_spacing(company_p, before=0, after=1, line=1.1)
    run = company_p.add_run(f"{company} | {dates}")
    set_run_font(run, size=10, color=MUTED)

    if meta:
        body_para(doc, meta, size=9.5, after=2, before=0, color=MUTED)
    if bullets:
        for b in bullets:
            bullet_para(doc, b)


def build_docx(path: Path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    name = doc.add_paragraph()
    set_paragraph_spacing(name, before=0, after=2, line=1.0)
    run = name.add_run("HASSAN AMIN")
    set_run_font(run, size=18, bold=True)

    role = doc.add_paragraph()
    set_paragraph_spacing(role, before=0, after=4, line=1.0)
    run = role.add_run(TARGET_TITLE)
    set_run_font(run, size=11, color=MUTED)

    for i, line in enumerate(CONTACT_LINES):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=1 if i == 0 else 2, line=1.1)
        run = p.add_run(line)
        set_run_font(run, size=9.5, color=MUTED)

    section_heading(doc, "Professional Summary")
    body_para(doc, SUMMARY, size=10, after=2, justify=True)

    section_heading(doc, "Skills")
    for label, value in SKILLS:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=2, after=2, line=1.15)
        run = p.add_run(f"{label}: ")
        set_run_font(run, size=10, bold=True)
        run = p.add_run(value)
        set_run_font(run, size=10)

    section_heading(doc, "Professional Experience")
    for job in EXPERIENCE:
        add_job_block(
            doc,
            job["title"],
            job["company"],
            job["dates"],
            meta=job.get("meta"),
            bullets=job["bullets"],
        )

    section_heading(doc, "Additional Experience")
    body_para(
        doc,
        "Freelance and consulting work completed alongside full-time roles.",
        size=9,
        after=3,
        color=MUTED,
    )
    for job in FREELANCE:
        add_job_block(
            doc,
            job["title"],
            job["company"],
            job["dates"],
            meta=job.get("meta"),
            bullets=job["bullets"],
        )

    section_heading(doc, "Education")
    for title, school in EDUCATION:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=3, after=0, line=1.1)
        run = p.add_run(title)
        set_run_font(run, size=10, bold=True)
        body_para(doc, school, size=10, after=2, color=MUTED)

    section_heading(doc, "Certifications")
    for cert in CERTIFICATIONS:
        bullet_para(doc, cert)

    section_heading(doc, "Languages")
    body_para(doc, "English: C2 | Arabic: Native", size=10, after=2)

    section_heading(doc, "Portfolio")
    body_para(doc, "https://hassanamin.net", size=10.5, after=2)

    doc.save(path)


def build_pdf(path: Path):
    doc = SimpleDocTemplate(
        str(path),
        pagesize=letter,
        leftMargin=0.7 * inch,
        rightMargin=0.7 * inch,
        topMargin=0.5 * inch,
        bottomMargin=0.5 * inch,
    )
    base = getSampleStyleSheet()["Normal"]
    name_s = ParagraphStyle(
        "Name", parent=base, fontName="Helvetica-Bold", fontSize=17,
        leading=20, textColor=HexColor(INK_HEX), spaceAfter=2, alignment=TA_LEFT,
    )
    role_s = ParagraphStyle(
        "RoleLine", parent=base, fontName="Helvetica", fontSize=10.5,
        leading=13, textColor=HexColor(MUTED_HEX), spaceAfter=4, alignment=TA_LEFT,
    )
    contact_s = ParagraphStyle(
        "ContactLine", parent=base, fontName="Helvetica", fontSize=9,
        leading=12, textColor=HexColor(MUTED_HEX), spaceAfter=1, alignment=TA_LEFT,
    )
    section_s = ParagraphStyle(
        "SectionHead", parent=base, fontName="Helvetica-Bold", fontSize=9.5,
        leading=12, textColor=HexColor(INK_HEX), spaceBefore=10, spaceAfter=2,
        alignment=TA_LEFT,
    )
    body_s = ParagraphStyle(
        "BodyText", parent=base, fontName="Helvetica", fontSize=9,
        leading=12, textColor=HexColor(INK_HEX), spaceAfter=4, alignment=TA_JUSTIFY,
    )
    muted_s = ParagraphStyle(
        "MutedText", parent=base, fontName="Helvetica", fontSize=8.5,
        leading=11, textColor=HexColor(MUTED_HEX), spaceAfter=2, alignment=TA_LEFT,
    )
    job_title_s = ParagraphStyle(
        "JobTitle", parent=base, fontName="Helvetica-Bold", fontSize=10,
        leading=12, textColor=HexColor(INK_HEX), spaceBefore=6, spaceAfter=0,
        alignment=TA_LEFT,
    )
    job_company_s = ParagraphStyle(
        "JobCompany", parent=base, fontName="Helvetica", fontSize=9,
        leading=11, textColor=HexColor(MUTED_HEX), spaceAfter=1, alignment=TA_LEFT,
    )
    bullet_s = ParagraphStyle(
        "BulletText", parent=base, fontName="Helvetica", fontSize=9,
        leading=12, textColor=HexColor(INK_HEX), leftIndent=10, spaceAfter=2,
        alignment=TA_LEFT,
    )
    skill_s = ParagraphStyle(
        "SkillLine", parent=base, fontName="Helvetica", fontSize=9,
        leading=12, textColor=HexColor(INK_HEX), spaceBefore=1, spaceAfter=2,
        alignment=TA_LEFT,
    )

    story = []
    story.append(Paragraph("HASSAN AMIN", name_s))
    story.append(Paragraph(TARGET_TITLE, role_s))
    for line in CONTACT_LINES:
        story.append(Paragraph(line.replace(" | ", " &nbsp;|&nbsp; "), contact_s))
    story.append(Spacer(1, 4))

    def section(title):
        story.append(Paragraph(title.upper(), section_s))
        story.append(
            HRFlowable(
                width="100%", thickness=0.75, color=HexColor(RULE_HEX),
                spaceBefore=0, spaceAfter=5,
            )
        )

    def job_block(job):
        story.append(Paragraph(job["title"], job_title_s))
        story.append(Paragraph(f"{job['company']} | {job['dates']}", job_company_s))
        if job.get("meta"):
            story.append(Paragraph(job["meta"], muted_s))
        for b in job["bullets"]:
            story.append(Paragraph(f"- {b}", bullet_s))

    section("Professional Summary")
    story.append(Paragraph(SUMMARY, body_s))

    section("Skills")
    for label, value in SKILLS:
        story.append(Paragraph(f"<b>{label}:</b> {value}", skill_s))

    section("Professional Experience")
    for job in EXPERIENCE:
        job_block(job)

    section("Additional Experience")
    story.append(
        Paragraph(
            "Freelance and consulting work completed alongside full-time roles.",
            muted_s,
        )
    )
    for job in FREELANCE:
        job_block(job)

    section("Education")
    for title, school in EDUCATION:
        story.append(Paragraph(f"<b>{title}</b>", skill_s))
        story.append(Paragraph(school, muted_s))

    section("Certifications")
    for cert in CERTIFICATIONS:
        story.append(Paragraph(f"- {cert}", bullet_s))

    section("Languages")
    story.append(Paragraph("English: C2 | Arabic: Native", skill_s))

    section("Portfolio")
    story.append(Paragraph("https://hassanamin.net", skill_s))

    doc.build(story)


def build_txt(path: Path):
    lines = [
        "HASSAN AMIN",
        TARGET_TITLE,
        *CONTACT_LINES,
        "",
        "PROFESSIONAL SUMMARY",
        SUMMARY,
        "",
        "SKILLS",
    ]
    for label, value in SKILLS:
        lines.append(f"{label}: {value}")
    lines.extend(["", "PROFESSIONAL EXPERIENCE", ""])
    for job in EXPERIENCE:
        lines.append(job["title"])
        lines.append(f"{job['company']} | {job['dates']}")
        if job.get("meta"):
            lines.append(job["meta"])
        for b in job["bullets"]:
            lines.append(f"- {b}")
        lines.append("")
    lines.append("ADDITIONAL EXPERIENCE")
    lines.append("Freelance and consulting work completed alongside full-time roles.")
    lines.append("")
    for job in FREELANCE:
        lines.append(job["title"])
        lines.append(f"{job['company']} | {job['dates']}")
        if job.get("meta"):
            lines.append(job["meta"])
        for b in job["bullets"]:
            lines.append(f"- {b}")
        lines.append("")
    lines.append("EDUCATION")
    for title, school in EDUCATION:
        lines.append(title)
        lines.append(school)
        lines.append("")
    lines.append("CERTIFICATIONS")
    for cert in CERTIFICATIONS:
        lines.append(f"- {cert}")
    lines.extend(["", "LANGUAGES", "English: C2 | Arabic: Native", "", "PORTFOLIO", "https://hassanamin.net"])
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    docx_path = OUT_DIR / f"{BASE}.docx"
    pdf_path = OUT_DIR / f"{BASE}.pdf"
    txt_path = OUT_DIR / f"{BASE}.txt"
    build_docx(docx_path)
    build_pdf(pdf_path)
    build_txt(txt_path)
    print(f"Wrote {docx_path}")
    print(f"Wrote {pdf_path}")
    print(f"Wrote {txt_path}")


if __name__ == "__main__":
    main()
