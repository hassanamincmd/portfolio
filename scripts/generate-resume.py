"""Generate ATS-optimized resume (DOCX + PDF + TXT) for Hassan Amin.

ATS hierarchy (optimized for parsing + keyword scoring):
  1. Contact / target title
  2. Professional Summary
  3. Skills              <- keywords early
  4. Professional Experience
  5. Additional Experience (freelance; kept separate from FT timeline)
  6. Education
  7. Certifications
  8. Portfolio           <- URL also in contact; detail at end
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.colors import HexColor
from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import HRFlowable, Paragraph, SimpleDocTemplate, Spacer

ROOT = Path(__file__).resolve().parents[1]
ASSETS = ROOT / "assets"
BASE = "Hassan-Amin-CV-Resume"

INK = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x5A, 0x5A, 0x5A)
INK_HEX = "#1A1A1A"
MUTED_HEX = "#5A5A5A"
RULE_HEX = "#C8C8C8"

TARGET_TITLE = "Senior Product Designer | UI/UX Designer"

PORTFOLIO_URL = "https://hassanamin.net"
PORTFOLIO_BLURB = "Personal portfolio and case studies."

CONTACT_LINES = [
    "Cairo, Egypt | +20 106 687 4777 | contact.hassan.amin@gmail.com",
    "LinkedIn: linkedin.com/in/hassan-mo-amin | Portfolio: https://hassanamin.net",
]

SUMMARY = (
    "UI/UX Manager and Product Designer with 5 years of experience across enterprise SaaS, "
    "agency products, and government-facing work in MENA and Asia. Managed a team of 4 at "
    "Caspian delivering 20+ digital products annually with 100% on-time delivery. Most recently "
    "Product Designer II at Procore (Quality and Safety), contributing to research tied to a "
    "25% increase in safety tools engagement. Applying for Senior Product Designer or Senior "
    "UI/UX Designer roles."
)

EXPERIENCE = [
    {
        "title": "Product Designer II",
        "company": "Procore Technologies",
        "dates": "May 2025 - August 2026",
        "meta": "Quality and Safety Design Team | Full-time",
        "bullets": [
            "Designed product UI for Quality and Safety tools on Procore's construction management platform.",
            "Contributed to research used to shape the Quality and Safety domain.",
            "Work contributed to a 25% increase in safety tools engagement.",
        ],
    },
    {
        "title": "UI/UX Manager (Senior)",
        "company": "Caspian Digital Solutions",
        "dates": "September 2022 - February 2025",
        "meta": "Full-time",
        "bullets": [
            "Managed and directed a UX design team of 4, delivering over 20 high-quality digital products annually.",
            "Ensured 100% on-time delivery with a 95% stakeholder approval rate.",
            "Designed high-fidelity prototypes and wireframes, increasing client satisfaction by 35%.",
            "Implemented quarterly reviews to optimize workflows, improving team efficiency by 20%.",
        ],
    },
    {
        "title": "UI/UX Designer",
        "company": "Caspian Digital Solutions",
        "dates": "May 2022 - November 2022",
        "meta": "Full-time",
        "bullets": [
            "Created detailed wireframes and user flows, improving cross-team clarity by 20%.",
            "Established design guidelines, enhancing brand consistency and reducing iteration time by 25%.",
            "Presented and defended designs, achieving a 90% approval rate in the first round.",
        ],
    },
    {
        "title": "UI/UX Designer / Product Designer",
        "company": "Spiritude",
        "dates": "December 2021 - May 2022",
        "meta": "Full-time",
        "bullets": [
            "Conducted user research and competitor analysis, driving a 20% improvement in product positioning.",
            "Created user personas and usage scenarios, increasing feature adoption by 30%.",
            "Delivered cohesive design systems, reducing inconsistencies by 40%.",
        ],
    },
]

# Side / concurrent work - kept separate from the full-time timeline.
FREELANCE = [
    {
        "title": "Senior Product Designer (Freelance)",
        "company": "Deloitte",
        "dates": "2025",
        "meta": "Freelance",
        "bullets": [
            "Designed UI components for the KSA Capital Market Authority design system.",
        ],
    },
    {
        "title": "Senior Product Designer / UX Strategy & Design Lead",
        "company": "Path2Live",
        "dates": "March 2024 - December 2024",
        "meta": "Contractor",
        "bullets": [
            "Led the UI/UX team in designing and delivering multiple projects, with handoff to development teams.",
            "Established workflows and structured the company's design team, improving team efficiency by 40%.",
            "Developed a comprehensive design system, improving collaboration between designers and developers.",
            "Presented design concepts to stakeholders, reducing decision-making time by 30%.",
            "Delivered responsive, cross-platform designs, reducing integration issues by 15%.",
        ],
    },
]

EDUCATION = [
    ("Diploma in UI/UX Design", "Information Technology Institute (ITI)"),
    ("B.A. in Geographic Information Systems", "King Marriott Academy"),
]

CERTIFICATIONS = [
    "Google UX Design Professional Certificate",
    "Digital Product Management (IMTF)",
    "Graphic Design Practice License (Egypt)",
    "PMP Training (AASTMT)",
]

SKILLS = [
    (
        "Design and Prototyping",
        "Figma, Framer, Adobe XD, Photoshop, Illustrator, Wireframing, Prototyping, "
        "Interaction Design, Responsive Design, Design Systems, Accessibility (WCAG), UX Writing",
    ),
    (
        "Research and Strategy",
        "User Research, Usability Testing, UserTesting, Synthesis, Data-Driven Design, "
        "Design Thinking, Agile",
    ),
    (
        "AI-Assisted Workflow",
        "Cursor, Claude, ChatGPT, Prompt Engineering, AI-Assisted Prototyping and Handoff",
    ),
    (
        "Collaboration",
        "Jira, Notion, Trello, Stakeholder Workshops, Design-Engineering Handoff",
    ),
]


# ---------------------------------------------------------------------------
# DOCX helpers
# ---------------------------------------------------------------------------

def set_run_font(run, size=10.5, bold=False, color=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color if color is not None else INK


def set_paragraph_spacing(p, before=0, after=0, line=1.15):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line


def add_bottom_border(paragraph, color="C8C8C8", size="12"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def section_heading(doc, text):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=12, after=5, line=1.0)
    run = p.add_run(text.upper())
    set_run_font(run, size=10, bold=True, color=INK)
    add_bottom_border(p)
    return p


def body_para(doc, text, size=10, after=6, before=0, justify=False, color=None):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=before, after=after, line=1.2)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_run_font(run, size=size, color=color)
    return p


def bullet_para(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_paragraph_spacing(p, before=0, after=2, line=1.2)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.first_line_indent = Inches(-0.15)
    if p.runs:
        p.runs[0].text = text
        set_run_font(p.runs[0], size=10)
    else:
        run = p.add_run(text)
        set_run_font(run, size=10)
    return p


def add_job_block(doc, title, company, dates, meta=None, bullets=None, line=None):
    """ATS-friendly job block: Title / Company | Dates / optional meta / bullets."""
    title_p = doc.add_paragraph()
    set_paragraph_spacing(title_p, before=9, after=0, line=1.1)
    run = title_p.add_run(title)
    set_run_font(run, size=10.5, bold=True)

    company_p = doc.add_paragraph()
    set_paragraph_spacing(company_p, before=0, after=1, line=1.1)
    run = company_p.add_run(f"{company} | {dates}")
    set_run_font(run, size=10, color=MUTED)

    if meta:
        body_para(doc, meta, size=9.5, after=2, before=0, color=MUTED)

    if bullets:
        for b in bullets:
            bullet_para(doc, b)
    if line:
        body_para(doc, line, size=10, after=2)


def add_skills_block(doc):
    for label, value in SKILLS:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=2, after=2, line=1.2)
        run = p.add_run(f"{label}: ")
        set_run_font(run, size=10, bold=True)
        run = p.add_run(value)
        set_run_font(run, size=10)


def build_docx(path: Path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    name = doc.add_paragraph()
    set_paragraph_spacing(name, before=0, after=2, line=1.0)
    run = name.add_run("HASSAN AMIN")
    set_run_font(run, size=20, bold=True)

    role = doc.add_paragraph()
    set_paragraph_spacing(role, before=0, after=4, line=1.0)
    run = role.add_run(TARGET_TITLE)
    set_run_font(run, size=11, color=MUTED)

    for i, line in enumerate(CONTACT_LINES):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=1 if i == 0 else 2, line=1.1)
        run = p.add_run(line)
        set_run_font(run, size=9.5, color=MUTED)

    # 1. Professional Summary
    section_heading(doc, "Professional Summary")
    body_para(doc, SUMMARY, size=10, after=2, justify=True)

    # 2. Skills (keywords early for ATS scoring)
    section_heading(doc, "Skills")
    add_skills_block(doc)

    # 3. Professional Experience
    section_heading(doc, "Professional Experience")
    for job in EXPERIENCE:
        add_job_block(
            doc,
            job["title"],
            job["company"],
            job["dates"],
            meta=job["meta"],
            bullets=job["bullets"],
        )

    # 4. Additional Experience (freelance - separate from FT timeline)
    section_heading(doc, "Additional Experience")
    body_para(
        doc,
        "Freelance and consulting work done alongside full-time roles.",
        size=9,
        after=3,
        color=MUTED,
    )
    for job in FREELANCE:
        add_job_block(
            doc,
            job["title"],
            job["company"],
            job["dates"],
            meta=job.get("meta"),
            bullets=job["bullets"],
        )

    # 5. Education
    section_heading(doc, "Education")
    for title, school in EDUCATION:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=3, after=0, line=1.15)
        run = p.add_run(title)
        set_run_font(run, size=10, bold=True)
        body_para(doc, school, size=10, after=2, color=MUTED)

    # 6. Certifications
    section_heading(doc, "Certifications")
    for cert in CERTIFICATIONS:
        bullet_para(doc, cert)

    # 7. Portfolio (URL already in contact; detail last)
    section_heading(doc, "Portfolio")
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=2, line=1.2)
    run = p.add_run(PORTFOLIO_URL)
    set_run_font(run, size=10.5, bold=True)
    body_para(doc, PORTFOLIO_BLURB, size=10, after=2, color=MUTED)

    doc.save(path)


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

def build_pdf(path: Path):
    doc = SimpleDocTemplate(
        str(path),
        pagesize=letter,
        leftMargin=0.7 * inch,
        rightMargin=0.7 * inch,
        topMargin=0.5 * inch,
        bottomMargin=0.5 * inch,
    )
    base = getSampleStyleSheet()["Normal"]

    name_s = ParagraphStyle(
        "Name", parent=base, fontName="Helvetica-Bold", fontSize=18,
        leading=22, textColor=HexColor(INK_HEX), spaceAfter=2, alignment=TA_LEFT,
    )
    role_s = ParagraphStyle(
        "RoleLine", parent=base, fontName="Helvetica", fontSize=10.5,
        leading=13, textColor=HexColor(MUTED_HEX), spaceAfter=4, alignment=TA_LEFT,
    )
    contact_s = ParagraphStyle(
        "ContactLine", parent=base, fontName="Helvetica", fontSize=9,
        leading=12, textColor=HexColor(MUTED_HEX), spaceAfter=1, alignment=TA_LEFT,
    )
    section_s = ParagraphStyle(
        "SectionHead", parent=base, fontName="Helvetica-Bold", fontSize=9.5,
        leading=12, textColor=HexColor(INK_HEX), spaceBefore=11, spaceAfter=2,
        alignment=TA_LEFT,
    )
    body_s = ParagraphStyle(
        "BodyText", parent=base, fontName="Helvetica", fontSize=9.5,
        leading=12.5, textColor=HexColor(INK_HEX), spaceAfter=4,
        alignment=TA_JUSTIFY,
    )
    muted_s = ParagraphStyle(
        "MutedText", parent=base, fontName="Helvetica", fontSize=9,
        leading=12, textColor=HexColor(MUTED_HEX), spaceAfter=2, alignment=TA_LEFT,
    )
    job_title_s = ParagraphStyle(
        "JobTitle", parent=base, fontName="Helvetica-Bold", fontSize=10,
        leading=13, textColor=HexColor(INK_HEX), spaceBefore=7, spaceAfter=0,
        alignment=TA_LEFT,
    )
    job_company_s = ParagraphStyle(
        "JobCompany", parent=base, fontName="Helvetica", fontSize=9.5,
        leading=12, textColor=HexColor(MUTED_HEX), spaceAfter=1, alignment=TA_LEFT,
    )
    bullet_s = ParagraphStyle(
        "BulletText", parent=base, fontName="Helvetica", fontSize=9.5,
        leading=12.5, textColor=HexColor(INK_HEX), leftIndent=10,
        spaceAfter=2, alignment=TA_LEFT,
    )
    skill_s = ParagraphStyle(
        "SkillLine", parent=base, fontName="Helvetica", fontSize=9.5,
        leading=12.5, textColor=HexColor(INK_HEX), spaceBefore=1, spaceAfter=2,
        alignment=TA_LEFT,
    )
    portfolio_url_s = ParagraphStyle(
        "PortfolioUrl", parent=base, fontName="Helvetica-Bold", fontSize=10,
        leading=13, textColor=HexColor(INK_HEX), spaceAfter=2, alignment=TA_LEFT,
    )

    story = []
    story.append(Paragraph("HASSAN AMIN", name_s))
    story.append(Paragraph(TARGET_TITLE, role_s))
    for line in CONTACT_LINES:
        story.append(Paragraph(line.replace(" | ", " &nbsp;|&nbsp; "), contact_s))
    story.append(Spacer(1, 4))

    def section(title):
        story.append(Paragraph(title.upper(), section_s))
        story.append(
            HRFlowable(
                width="100%", thickness=0.75, color=HexColor(RULE_HEX),
                spaceBefore=0, spaceAfter=5,
            )
        )

    def job_block(title, company, dates, meta=None, bullets=None, line=None):
        story.append(Paragraph(title, job_title_s))
        story.append(Paragraph(f"{company} | {dates}", job_company_s))
        if meta:
            story.append(Paragraph(meta, muted_s))
        if bullets:
            for b in bullets:
                story.append(Paragraph(f"- {b}", bullet_s))
        if line:
            story.append(Paragraph(f"- {line}", bullet_s))

    section("Professional Summary")
    story.append(Paragraph(SUMMARY, body_s))

    section("Skills")
    for label, value in SKILLS:
        story.append(Paragraph(f"<b>{label}:</b> {value}", skill_s))

    section("Professional Experience")
    for job in EXPERIENCE:
        job_block(
            job["title"],
            job["company"],
            job["dates"],
            meta=job["meta"],
            bullets=job["bullets"],
        )

    section("Additional Experience")
    story.append(
        Paragraph(
            "Freelance and consulting work done alongside full-time roles.",
            muted_s,
        )
    )
    for job in FREELANCE:
        job_block(
            job["title"],
            job["company"],
            job["dates"],
            meta=job.get("meta"),
            bullets=job["bullets"],
        )

    section("Education")
    for title, school in EDUCATION:
        story.append(Paragraph(f"<b>{title}</b>", skill_s))
        story.append(Paragraph(school, muted_s))

    section("Certifications")
    for cert in CERTIFICATIONS:
        story.append(Paragraph(f"- {cert}", bullet_s))

    section("Portfolio")
    story.append(Paragraph(PORTFOLIO_URL, portfolio_url_s))
    story.append(Paragraph(PORTFOLIO_BLURB, muted_s))

    doc.build(story)


# ---------------------------------------------------------------------------
# TXT
# ---------------------------------------------------------------------------

def build_txt(path: Path):
    lines = [
        "HASSAN AMIN",
        TARGET_TITLE,
        *CONTACT_LINES,
        "",
        "PROFESSIONAL SUMMARY",
        SUMMARY,
        "",
        "SKILLS",
    ]
    for label, value in SKILLS:
        lines.append(f"{label}: {value}")
    lines.extend(["", "PROFESSIONAL EXPERIENCE", ""])
    for job in EXPERIENCE:
        lines.append(job["title"])
        lines.append(f"{job['company']} | {job['dates']}")
        if job["meta"]:
            lines.append(job["meta"])
        for b in job["bullets"]:
            lines.append(f"- {b}")
        lines.append("")
    lines.append("ADDITIONAL EXPERIENCE")
    lines.append("Freelance and consulting work done alongside full-time roles.")
    lines.append("")
    for job in FREELANCE:
        lines.append(job["title"])
        lines.append(f"{job['company']} | {job['dates']}")
        if job.get("meta"):
            lines.append(job["meta"])
        for b in job["bullets"]:
            lines.append(f"- {b}")
        lines.append("")
    lines.append("EDUCATION")
    for title, school in EDUCATION:
        lines.append(title)
        lines.append(school)
        lines.append("")
    lines.append("CERTIFICATIONS")
    for cert in CERTIFICATIONS:
        lines.append(f"- {cert}")
    lines.extend(["", "PORTFOLIO", PORTFOLIO_URL, PORTFOLIO_BLURB])
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main():
    ASSETS.mkdir(parents=True, exist_ok=True)
    docx_path = ASSETS / f"{BASE}.docx"
    pdf_path = ASSETS / f"{BASE}.pdf"
    txt_path = ASSETS / f"{BASE}.txt"
    build_docx(docx_path)
    build_pdf(pdf_path)
    build_txt(txt_path)
    print(f"Wrote {docx_path}")
    print(f"Wrote {pdf_path}")
    print(f"Wrote {txt_path}")


if __name__ == "__main__":
    main()
